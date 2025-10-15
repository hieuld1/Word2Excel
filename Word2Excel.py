#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Copy dữ liệu từ bảng trong Word sang vùng Excel (hàng H1..H2, cột C1..C2),
bật Wrap Text, canh Top và kẻ Border outline + inside.

Cài đặt:
  pip install python-docx openpyxl
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.utils import column_index_from_string, get_column_letter
from openpyxl.styles import Alignment, Border, Side


def parse_col_arg(x: str) -> int:
    x = str(x).strip()
    if x.isdigit():
        val = int(x)
        if val < 1:
            raise ValueError("Chỉ số cột phải >= 1")
        return val
    return column_index_from_string(x.upper())


def read_table_from_word(word_path: Path, table_index: int = 0) -> list[list[str]]:
    doc = Document(word_path)
    if table_index < 0 or table_index >= len(doc.tables):
        raise IndexError(f"File Word có {len(doc.tables)} bảng, nhưng yêu cầu bảng index={table_index}")
    tbl = doc.tables[table_index]

    data: list[list[str]] = []
    for row in tbl.rows:
        row_vals = []
        for cell in row.cells:
            txt = " ".join(cell.text.replace("\xa0", " ").split())
            row_vals.append(txt)
        data.append(row_vals)
    return data


def ensure_workbook(path: Path):
    if path.exists():
        try:
            return load_workbook(path)
        except Exception:
            return Workbook()
    else:
        return Workbook()


def main():
    ap = argparse.ArgumentParser(description="Copy bảng Word -> vùng Excel (Wrap Text + Border)")
    ap.add_argument("--word-in", required=True, help="Đường dẫn file Word (.docx)")
    ap.add_argument("--table-index", type=int, default=0, help="Index bảng trong Word (0 = bảng đầu tiên)")
    ap.add_argument("--excel-out", required=True, help="File Excel đầu ra (.xlsx)")
    ap.add_argument("--sheet", default=None, help="Tên sheet (mặc định: active)")
    ap.add_argument("--row-start", type=int, required=True, help="H1: hàng bắt đầu (>=1)")
    ap.add_argument("--row-end", type=int, required=True, help="H2: hàng kết thúc (>=H1)")
    ap.add_argument("--col-start", required=True, help="C1: cột bắt đầu (vd: C hoặc 3)")
    ap.add_argument("--col-end", required=True, help="C2: cột kết thúc (vd: H hoặc 8)")
    args = ap.parse_args()

    word_path = Path(args.word_in)
    if not word_path.exists():
        print(f"❌ Không tìm thấy file Word: {word_path}")
        sys.exit(1)

    try:
        c1 = parse_col_arg(args.col_start)
        c2 = parse_col_arg(args.col_end)
    except Exception as e:
        print(f"❌ Lỗi tham số cột: {e}")
        sys.exit(1)

    r1, r2 = int(args.row_start), int(args.row_end)
    if r1 < 1 or r2 < r1 or c1 < 1 or c2 < c1:
        print("❌ Phạm vi hàng/cột không hợp lệ.")
        sys.exit(1)

    # Đọc bảng Word
    try:
        table_data = read_table_from_word(word_path, args.table_index)
    except Exception as e:
        print(f"❌ Lỗi đọc Word: {e}")
        sys.exit(1)

    rows_word = len(table_data)
    cols_word = max(len(r) for r in table_data)
    print(f"📖 Bảng Word[{args.table_index}]: {rows_word} x {cols_word}")

    # Mở hoặc tạo Excel
    wb = ensure_workbook(Path(args.excel_out))
    ws = wb[args.sheet] if args.sheet and args.sheet in wb.sheetnames else wb.active

    nrows_target = r2 - r1 + 1
    ncols_target = c2 - c1 + 1
    rows_to_write = min(rows_word, nrows_target)
    cols_to_write = min(cols_word, ncols_target)

    # Style cơ bản
    align = Alignment(wrap_text=True, vertical="top")
    thin = Side(border_style="thin", color="000000")
    border_style = Border(left=thin, right=thin, top=thin, bottom=thin)

    print(f"✍️  Ghi dữ liệu vào vùng hàng {r1}-{r2}, cột {c1}-{c2}...")

    # Ghi dữ liệu và style
    for i in range(rows_to_write):
        for j in range(cols_to_write):
            val = ""
            if j < len(table_data[i]):
                val = table_data[i][j]
            cell = ws.cell(row=r1 + i, column=c1 + j, value=val)
            cell.alignment = align
            cell.border = border_style

    # Kẻ border cho toàn vùng (bao gồm phần trống)
    for rr in range(r1, r2 + 1):
        for cc in range(c1, c2 + 1):
            cell = ws.cell(row=rr, column=cc)
            if cell.border is None or cell.border.left is None:
                cell.border = border_style
            cell.alignment = align

    # Lưu file
    rng = f"{get_column_letter(c1)}{r1}:{get_column_letter(c2)}{r2}"
    wb.save(args.excel_out)
    print(f"✅ Đã ghi bảng, Wrap Text và Border cho vùng {rng} -> {Path(args.excel_out).resolve()}")


if __name__ == "__main__":
    main()
