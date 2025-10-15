#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Chương trình: Copy dữ liệu vùng Excel (hàng H1..H2, cột C1..C2) sang bảng trong Word.

Cài đặt:
  pip install openpyxl python-docx

Ví dụ chạy:
  python excel_to_word_table.py --excel input.xlsx --sheet Sheet1 \
      --row-start 2 --row-end 10 --col-start C --col-end H --word-out output.docx
"""

import argparse
import sys
from pathlib import Path

from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def parse_col_arg(x: str) -> int:
    """
    Chấp nhận 'C' hoặc '3' -> trả về chỉ số cột dạng số (1-based).
    """
    x = str(x).strip()
    if x.isdigit():
        val = int(x)
        if val < 1:
            raise ValueError("Chỉ số cột phải >= 1")
        return val
    # Cột dạng chữ
    return column_index_from_string(x.upper())

def main():
    ap = argparse.ArgumentParser(description="Copy vùng Excel -> bảng Word")
    ap.add_argument("--excel", required=True, help="Đường dẫn file Excel (vd: input.xlsx)")
    ap.add_argument("--sheet", default=None, help="Tên sheet (mặc định: sheet active)")
    ap.add_argument("--row-start", type=int, required=True, help="Hàng bắt đầu (H1)")
    ap.add_argument("--row-end", type=int, required=True, help="Hàng kết thúc (H2)")
    ap.add_argument("--col-start", required=True, help="Cột bắt đầu (C1) - chấp nhận 'C' hoặc '3'")
    ap.add_argument("--col-end", required=True, help="Cột kết thúc (C2) - chấp nhận 'H' hoặc '8'")
    ap.add_argument("--word-out", default="output.docx", help="File Word đầu ra (mặc định: output.docx)")
    args = ap.parse_args()

    excel_path = Path(args.excel)
    if not excel_path.exists():
        print(f"❌ Không tìm thấy file Excel: {excel_path}")
        sys.exit(1)

    # Chuyển cột sang số
    try:
        c1 = parse_col_arg(args.col_start)
        c2 = parse_col_arg(args.col_end)
    except Exception as e:
        print(f"❌ Lỗi cột C1/C2: {e}")
        sys.exit(1)

    r1, r2 = int(args.row_start), int(args.row_end)
    if r1 < 1 or r2 < 1:
        print("❌ Hàng phải >= 1")
        sys.exit(1)
    if c1 > c2 or r1 > r2:
        print("❌ Phạm vi không hợp lệ: cần C1<=C2 và H1<=H2")
        sys.exit(1)

    # Đọc Excel
    try:
        wb = load_workbook(excel_path, data_only=True)
        ws = wb[args.sheet] if args.sheet else wb.active
    except KeyError:
        print(f"❌ Không tìm thấy sheet: {args.sheet}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Lỗi mở Excel: {e}")
        sys.exit(1)

    nrows = r2 - r1 + 1
    ncols = c2 - c1 + 1
    print(f"📖 Đọc vùng: hàng {r1}..{r2}, cột {c1}..{c2} (kích thước {nrows}x{ncols})")

    # Lấy dữ liệu vùng
    data = []
    for r in range(r1, r2 + 1):
        row_vals = []
        for c in range(c1, c2 + 1):
            val = ws.cell(row=r, column=c).value
            row_vals.append("" if val is None else str(val))
        data.append(row_vals)

    # Tạo Word, chèn bảng
    try:
        doc = Document()

        # (Tuỳ chọn) đặt font mặc định
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)

        doc.add_heading("Dữ liệu từ Excel", level=1)
        tbl = doc.add_table(rows=nrows, cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        for i in range(nrows):
            for j in range(ncols):
                cell = tbl.cell(i, j)
                cell.text = data[i][j]

        out_path = Path(args.word_out)
        doc.save(out_path)
        print(f"✅ Đã ghi bảng vào file Word: {out_path.resolve()}")
    except Exception as e:
        print(f"❌ Lỗi ghi Word: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
